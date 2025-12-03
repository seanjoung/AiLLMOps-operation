#!/usr/bin/env python3
"""
Report Generator Module
CSV, DOCX 형식 보고서 생성
"""

import csv
import os
from datetime import datetime
from typing import List, Dict, Any
from dataclasses import dataclass

# DOCX 생성용
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class ReportConfig:
    report_type: str = "weekly"  # weekly or monthly
    company_name: str = "Company"
    team_name: str = "Infrastructure Team"
    output_dir: str = "./output"


class ReportGenerator:
    def __init__(self, config: ReportConfig = None):
        self.config = config or ReportConfig()
        os.makedirs(self.config.output_dir, exist_ok=True)
        
    def _get_report_title(self) -> str:
        now = datetime.now()
        if self.config.report_type == "weekly":
            week_num = now.isocalendar()[1]
            return f"{now.year}년 {week_num}주차 인프라 정기점검 보고서"
        else:
            return f"{now.year}년 {now.month}월 인프라 정기점검 보고서"
    
    def _get_filename_prefix(self) -> str:
        now = datetime.now()
        if self.config.report_type == "weekly":
            week_num = now.isocalendar()[1]
            return f"infra_check_{now.year}_W{week_num:02d}"
        else:
            return f"infra_check_{now.year}_{now.month:02d}"
    
    def generate_csv(self, results: List[Dict], summary: Dict) -> str:
        """CSV 보고서 생성"""
        filename = f"{self._get_filename_prefix()}.csv"
        filepath = os.path.join(self.config.output_dir, filename)
        
        with open(filepath, 'w', newline='', encoding='utf-8-sig') as f:
            # 헤더 정보 작성
            f.write(f"# {self._get_report_title()}\n")
            f.write(f"# 생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"# 회사: {self.config.company_name}\n")
            f.write(f"# 담당팀: {self.config.team_name}\n")
            f.write(f"# 총 점검항목: {summary.get('total', 0)}개\n")
            f.write(f"# 정상: {summary.get('ok', 0)} / 경고: {summary.get('warning', 0)} / 위험: {summary.get('critical', 0)} / 확인불가: {summary.get('unknown', 0)}\n")
            f.write("\n")
            
            # 데이터 작성
            if results:
                writer = csv.DictWriter(f, fieldnames=results[0].keys())
                writer.writeheader()
                writer.writerows(results)
        
        return filepath
    
    def generate_docx(self, results: List[Dict], summary: Dict) -> str:
        """DOCX 보고서 생성"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 라이브러리가 설치되지 않았습니다. pip install python-docx")
        
        filename = f"{self._get_filename_prefix()}.docx"
        filepath = os.path.join(self.config.output_dir, filename)
        
        doc = Document()
        
        # 문서 스타일 설정
        style = doc.styles['Normal']
        style.font.name = 'Malgun Gothic'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        
        # 제목
        title = doc.add_heading(self._get_report_title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 보고서 정보
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").bold = False
        info_para.add_run(f"회사: {self.config.company_name} | 담당팀: {self.config.team_name}")
        
        doc.add_paragraph()
        
        # 요약 섹션
        doc.add_heading('1. 점검 결과 요약', level=1)
        
        summary_table = doc.add_table(rows=2, cols=5)
        summary_table.style = 'Table Grid'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 헤더
        hdr_cells = summary_table.rows[0].cells
        headers = ['총 점검항목', '정상', '경고', '위험', '확인불가']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in hdr_cells[i].paragraphs[0].runs:
                run.bold = True
        
        # 데이터
        data_cells = summary_table.rows[1].cells
        data = [
            str(summary.get('total', 0)),
            str(summary.get('ok', 0)),
            str(summary.get('warning', 0)),
            str(summary.get('critical', 0)),
            str(summary.get('unknown', 0))
        ]
        colors = [None, RGBColor(0, 128, 0), RGBColor(255, 165, 0), RGBColor(255, 0, 0), RGBColor(128, 128, 128)]
        
        for i, (value, color) in enumerate(zip(data, colors)):
            data_cells[i].text = value
            data_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if color:
                for run in data_cells[i].paragraphs[0].runs:
                    run.font.color.rgb = color
                    run.bold = True
        
        doc.add_paragraph()
        
        # 카테고리별 요약
        doc.add_heading('2. 카테고리별 결과', level=1)
        
        by_cat = summary.get('by_category', {})
        cat_table = doc.add_table(rows=len(by_cat) + 1, cols=5)
        cat_table.style = 'Table Grid'
        
        cat_headers = ['카테고리', '정상', '경고', '위험', '확인불가']
        for i, h in enumerate(cat_headers):
            cat_table.rows[0].cells[i].text = h
            cat_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cat_table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        
        for row_idx, (cat_name, cat_data) in enumerate(by_cat.items(), start=1):
            cat_table.rows[row_idx].cells[0].text = cat_name
            cat_table.rows[row_idx].cells[1].text = str(cat_data.get('ok', 0))
            cat_table.rows[row_idx].cells[2].text = str(cat_data.get('warning', 0))
            cat_table.rows[row_idx].cells[3].text = str(cat_data.get('critical', 0))
            cat_table.rows[row_idx].cells[4].text = str(cat_data.get('unknown', 0))
            for cell in cat_table.rows[row_idx].cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 상세 결과
        doc.add_heading('3. 상세 점검 결과', level=1)
        
        # 카테고리별로 그룹화
        categories = {'OS': [], 'Kubernetes': [], 'Services': []}
        for r in results:
            cat = r.get('카테고리', 'Unknown')
            if cat in categories:
                categories[cat].append(r)
        
        for cat_name, cat_results in categories.items():
            if not cat_results:
                continue
                
            doc.add_heading(f'3.{list(categories.keys()).index(cat_name)+1} {cat_name} 점검', level=2)
            
            # 테이블 생성
            table = doc.add_table(rows=len(cat_results) + 1, cols=5)
            table.style = 'Table Grid'
            
            # 헤더
            headers = ['점검ID', '점검항목', '상태', '측정값', '결과메시지']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.bold = True
            
            # 데이터
            for row_idx, r in enumerate(cat_results, start=1):
                table.rows[row_idx].cells[0].text = r.get('점검ID', '')
                table.rows[row_idx].cells[1].text = r.get('점검항목', '')
                
                status = r.get('상태', '')
                table.rows[row_idx].cells[2].text = status
                # 상태별 색상
                status_cell = table.rows[row_idx].cells[2]
                for run in status_cell.paragraphs[0].runs:
                    if status == '정상':
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif status == '경고':
                        run.font.color.rgb = RGBColor(255, 165, 0)
                    elif status == '위험':
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    run.bold = True
                
                # 측정값 (너무 길면 자르기)
                value = r.get('측정값', '')
                if len(value) > 50:
                    value = value[:50] + "..."
                table.rows[row_idx].cells[3].text = value
                table.rows[row_idx].cells[4].text = r.get('결과메시지', '')
            
            doc.add_paragraph()
        
        # 문제 항목 요약
        issues = [r for r in results if r.get('상태') in ['경고', '위험']]
        if issues:
            doc.add_heading('4. 조치 필요 항목', level=1)
            for issue in issues:
                para = doc.add_paragraph()
                status = issue.get('상태', '')
                icon = "⚠️" if status == '경고' else "❌"
                run = para.add_run(f"{icon} [{issue.get('점검ID')}] {issue.get('점검항목')}")
                run.bold = True
                para.add_run(f"\n   - 상태: {status}")
                para.add_run(f"\n   - 내용: {issue.get('결과메시지', '')}")
                para.add_run(f"\n   - 설명: {issue.get('설명', '')}")
        
        # 서명란
        doc.add_paragraph()
        doc.add_paragraph()
        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_para.add_run("점검자: ________________")
        sign_para.add_run("\n\n")
        sign_para.add_run("검토자: ________________")
        
        doc.save(filepath)
        return filepath


def generate_reports(results: List[Dict], summary: Dict, config: ReportConfig = None) -> Dict[str, str]:
    """CSV와 DOCX 보고서 모두 생성"""
    generator = ReportGenerator(config)
    
    generated = {}
    
    # CSV 생성
    csv_path = generator.generate_csv(results, summary)
    generated['csv'] = csv_path
    
    # DOCX 생성
    if DOCX_AVAILABLE:
        docx_path = generator.generate_docx(results, summary)
        generated['docx'] = docx_path
    
    return generated


if __name__ == "__main__":
    # 테스트용 더미 데이터
    test_results = [
        {'점검ID': 'OS-001', '점검항목': '디스크 사용량', '카테고리': 'OS', 
         '설명': '테스트', '상태': '정상', '측정값': '45%', '임계치': '80%', 
         '결과메시지': '정상 범위 내', '점검시간': datetime.now().isoformat()},
    ]
    test_summary = {'total': 1, 'ok': 1, 'warning': 0, 'critical': 0, 'unknown': 0,
                    'by_category': {'OS': {'ok': 1, 'warning': 0, 'critical': 0, 'unknown': 0}}}
    
    config = ReportConfig(company_name="Test Corp", team_name="DevOps")
    paths = generate_reports(test_results, test_summary, config)
    print(f"Generated reports: {paths}")
